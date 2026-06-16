# -*- coding: utf-8 -*-
"""
周报生成器 — Weekly Report Generator
基于用户最终版第二十六次汇报的实际格式

【使用方法】
  1. 在本文件末尾的 ─── 每周内容填写区 ─── 中修改 content 字典
  2. 终端运行: python create_report.py
  3. 自动生成: {yyMMdd}-栾歆瑶第N次汇报.docx

【格式说明】
  - 标题: 黑体 18pt 加粗居中，前后段距 8pt
  - 章节标题: 宋体小四（12pt）加粗，首行缩进 24pt，1.5 倍行距
  - 正文: 宋体小四（12pt），首行缩进 24pt，1.5 倍行距
  - 子标题: 编号非加粗+标题加粗+内容普通，如"（1）安装[bold]与配置[bold]…"
  - 文献信息: 编号非加粗+标签加粗，如"（1）关键词：[bold]值[normal]"
  - 英文论文名后自动加中文译名行（纯文本无前缀）
  - 无空白行
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import re
from datetime import datetime

# ═══════════════════════════════════════════════════════════
#  格式常量
# ═══════════════════════════════════════════════════════════
STUDENT_NAME = "栾歆瑶"
FIRST_LINE_INDENT = Pt(24)
LINE_SPACING = 1.5
TITLE_SIZE = 18
BODY_SIZE = 12
TITLE_FONT = '黑体'
BODY_FONT = '宋体'


# ═══════════════════════════════════════════════════════════
#  字体 & 段落工具函数
# ═══════════════════════════════════════════════════════════

def _set_run(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _new_para(doc, indent=True, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    if indent:
        pf.first_line_indent = FIRST_LINE_INDENT
    if align:
        p.alignment = align
    return p


# ═══════════════════════════════════════════════════════════
#  文档构建函数
# ═══════════════════════════════════════════════════════════

def add_title(doc, text='周报总结'):
    """居中标题（黑体 18pt 加粗）"""
    p = _new_para(doc, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    _set_run(run, TITLE_FONT, TITLE_SIZE, bold=True)


def add_body(doc, text):
    """正文段落（宋体 12pt）"""
    p = _new_para(doc)
    run = p.add_run(text)
    _set_run(run, BODY_FONT, BODY_SIZE)


def add_section_title(doc, text):
    """章节标题（宋体 12pt 加粗）"""
    p = _new_para(doc)
    run = p.add_run(text)
    _set_run(run, BODY_FONT, BODY_SIZE, bold=True)


def add_mixed(doc, bold_part, normal_part='', indent=True):
    """加粗+普通混合"""
    p = _new_para(doc, indent=indent)
    run = p.add_run(bold_part)
    _set_run(run, BODY_FONT, BODY_SIZE, bold=True)
    if normal_part:
        run = p.add_run(normal_part)
        _set_run(run, BODY_FONT, BODY_SIZE)


def add_num_mixed(doc, num_part, bold_part, normal_part='', indent=True):
    """编号非加粗 + 标题加粗 + 内容普通
       如 （1）关键词：[bold]value[normal]
    """
    p = _new_para(doc, indent=indent)
    run = p.add_run(num_part)                 # 编号如（1）非加粗
    _set_run(run, BODY_FONT, BODY_SIZE, bold=False)
    run = p.add_run(bold_part)                # 标签如关键词：加粗
    _set_run(run, BODY_FONT, BODY_SIZE, bold=True)
    if normal_part:
        run = p.add_run(normal_part)           # 内容非加粗
        _set_run(run, BODY_FONT, BODY_SIZE)


def add_sub_item(doc, num, title, content='', indent=True):
    """子项：编号非加粗 + 标题加粗 + 内容普通（用于学习工作/项目工作）
       如 （1）下载[bold]与安装[bold]流程 内容内容
    """
    p = _new_para(doc, indent=indent)
    run = p.add_run(num)       # 编号如（1）非加粗
    _set_run(run, BODY_FONT, BODY_SIZE, bold=False)
    run = p.add_run(title)      # 标题如"下载与安装流程"加粗
    _set_run(run, BODY_FONT, BODY_SIZE, bold=True)
    if content:
        run = p.add_run(content)
        _set_run(run, BODY_FONT, BODY_SIZE)


# ═══════════════════════════════════════════════════════════
#  核心生成函数
# ═══════════════════════════════════════════════════════════

CHINESE_NUMS = ['一', '二', '三', '四', '五', '六', '七', '八']


def _get_next_week_num(output_dir=None):
    """从现有文件推断下一次汇报编号"""
    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(__file__))
    max_num = 26
    for f in os.listdir(output_dir):
        if f.endswith('.docx') and '栾歆瑶' in f:
            m = re.search(r'第(\d+)次', f)
            if m:
                n = int(m.group(1))
                if n > max_num:
                    max_num = n
    return max_num + 1


def generate_report(content: dict) -> str:
    """
    content 结构:
    {
        "summary": "一句话概要（逗号分隔）",       # 用于 "本周的学习汇报内容如下：①xx②xx"
        "project_works": [                          # 项目工作任务
            {
                "title": "项目标题",
                "items": [                          # 每个item是一段正文（无编号）
                    "内容1",
                    "内容2",
                ]
            },
        ],
        "study_works": [                            # 学习工作
            {
                "title": "学习标题",
                "items": [
                    ("（1）", "子标题加粗", "内容"),
                ],
            },
        ],
        "paper": {                                  # 文献阅读
            "title": "论文主题（用于章节名）",
            "title_en": "英文论文名称（纯文本）",
            "title_cn": "中文译名（纯文本）",
            "journal": "发表期刊（纯文本）",
            "year": "年份信息（纯文本）",
            "author": "作者信息（纯文本）",
            "numbered_items": [                     # 编号信息项
                ("关键词：", "值"),
                ("研究背景：", "值"),
                ("研究方法：", "值"),
                ("模型介绍：", "值"),
                ("实验设计：", "值"),
            ],
            "conclusion": [                         # 研究结论（列表，每项一段）
                "结论1",
                "结论2",
                ...
            ],
            "harvest": "（7）收获内容"              # 取代"对毕业论文的启发"
        }
    }
    """
    doc = Document()

    # 文档默认字体
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    # ── 1. 标题 ──
    add_title(doc)

    # ── 2. 概要 ──
    if content.get("summary"):
        add_body(doc, f'本周的学习汇报内容如下：{content["summary"]}。')

    # ── 3. 项目工作任务 ──
    sec_idx = 0
    for work in content.get("project_works", []):
        add_section_title(doc, f'{CHINESE_NUMS[sec_idx]}、{work["title"]}')
        sec_idx += 1
        if work.get("intro"):
            add_body(doc, work["intro"])
        for item in work.get("items", []):
            add_body(doc, item)

    # ── 4. 学习工作 ──
    for study in content.get("study_works", []):
        add_section_title(doc, f'{CHINESE_NUMS[sec_idx]}、{study["title"]}')
        sec_idx += 1
        if study.get("extra_body"):
            add_body(doc, study["extra_body"])
        for item in study.get("items", []):
            if len(item) == 3:
                num, sub_title, detail = item
                add_sub_item(doc, num, sub_title, detail)
            else:
                add_body(doc, item)

    # ── 5. 文献阅读 ──
    paper = content.get("paper")
    if paper:
        add_section_title(doc, f'{CHINESE_NUMS[sec_idx]}、文献阅读——{paper["title"]}')

        # 论文名称（纯文本，无前缀）
        if paper.get("title_en"):
            add_body(doc, paper["title_en"])
        # 中文译名（纯文本，无前缀）
        if paper.get("title_cn"):
            add_body(doc, paper["title_cn"])
        # 发表期刊
        if paper.get("journal"):
            add_body(doc, paper["journal"])
        # 年份
        if paper.get("year"):
            add_body(doc, paper["year"])
        # 作者信息
        if paper.get("author"):
            add_body(doc, paper["author"])

        # 编号信息项（1）~（5）: （1）关键词：值  [编号非加粗+标签加粗+值普通]
        nums_labels = ['（1）', '（2）', '（3）', '（4）', '（5）', '（6）', '（7）']
        for idx, (label, value) in enumerate(paper.get("numbered_items", [])):
            add_num_mixed(doc, nums_labels[idx], label, value)

        # 研究结论（6）：分条列出
        if paper.get("conclusion"):
            add_num_mixed(doc, '（6）', '研究结论：', '')
            for conclusion in paper["conclusion"]:
                if conclusion.startswith('①') or conclusion.startswith('②') or conclusion.startswith('③'):
                    add_body(doc, conclusion)
                else:
                    add_body(doc, conclusion)

        # 收获（7）取代"对毕业论文的启发"
        if paper.get("harvest"):
            p = _new_para(doc)
            run = p.add_run('（7）')
            _set_run(run, BODY_FONT, BODY_SIZE, bold=False)
            run = p.add_run('收获')
            _set_run(run, BODY_FONT, BODY_SIZE, bold=True)

    # ── 保存文件 ──
    today = datetime.now()
    date_str = today.strftime('%y%m%d')
    output_dir = os.path.dirname(os.path.abspath(__file__))
    week_num = _get_next_week_num(output_dir)

    filename = f'{date_str}-{STUDENT_NAME}第{week_num}次汇报.docx'
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f'周报已生成: {filename}')
    return output_path


# ═══════════════════════════════════════════════════════════
#  示例运行
# ═══════════════════════════════════════════════════════════
if __name__ == '__main__':

    content = {
        "summary": "①运筹说公众号推文修改 ②Claude Code学习 ③文献阅读",

        # ── 项目工作任务 ──
        "project_works": [
            {
                "title": "运筹说公众号推文修改",
                "intro": "本周根据老师提出的修改意见，对运筹说公众号推文的内容与排版进行了系统性修改，主要包括内容精简、可视化呈现、风格统一等方面的多轮优化，具体修改内容如下：",
                "items": [
                    '①大模型介绍内容精简：老师指出推文中大模型介绍部分文字过多，内容冗长。在修改中，对2.1节「主流大模型的格局」进行了大幅精简，修改为「模型名称+一句话核心优势+API价格」的紧凑格式，以半行精简表述替代大段文字。',
                    "②图表表现形式优化：将原表2模型参数规模演进改为演进图，采用时间轴+数据标签展示；将表4、表7改为信息图形式；在大模型格局部分优化了段落格式，参考了苹果发布会等视觉语言。",
                    "③整体风格调整：老师强调文章需要做出「大气磅礴」的感觉。在修改中，改为分节、分条的精简格式，减少了大段文字的堆砌。优化了对比部分的排版形式。",
                    "④表4改为演进图格式：将表4重新设计为时间轴演进图，展示2018-2026年训练数据从GB到EB级别的指数级增长过程，标注了BERT、GPT-3、Llama 3等关键里程碑。",
                    "⑤图表整体大气风格优化：采用深蓝色为主色调、金色为强调色的高端配色方案，去除冗余边框，采用极简主义设计原则。",
                ],
            },
        ],

        # ── 学习工作 ──
        "study_works": [
            {
                "title": "Claude Code人工智能编程助手的学习与使用",
                "items": [
                    ("（1）", "下载与安装流程",
                     "Claude Code是Anthropic公司推出的官方CLI工具，专为软件工程任务设计，能够在终端中直接与Claude进行交互，辅助完成代码编写、调试、重构、代码审查等开发工作。本周完成了Claude Code的完整安装与配置流程。"),
                    ("（2）", "基本原理与工作机制",
                     "Claude Code通过自然语言理解与代码生成能力辅助完成编程任务。其核心工作机制包括上下文感知能力、工具调用机制、任务规划与追踪功能、持久化记忆系统以及权限安全机制。"),
                    ("（3）", "Skill机制的学习与应用",
                     "安装完成后重点学习了Claude Code的Skill机制。具体实践包括：安装和使用了deep-research等文献综述相关的Skill。同时根据大模型、大算力等基础制作了ppt修改的相关skill。"),
                ],
            },
        ],

        # ── 文献阅读 ──
        "paper": {
            "title": "成品油配送领域的相关研究论文阅读与总结",
            "title_en": "Economical-traveling-distance-based fleet composition with fuel costs: An application in petrol distribution",
            "title_cn": "基于经济行驶距离的燃油成本车队组成：在成品油配送中的应用",
            "journal": "发表期刊：Transportation Research Part E",
            "year": "发表年份：2021年3月",
            "author": "作者信息：孙丽君（大连理工大学经济管理学院）、张源凯（浙江大学管理学院）、胡祥培（大连理工大学经济管理学院）",

            "numbered_items": [
                ("关键词：", "Fleet composition with fuel costs; Fleet size and mix; Economical traveling distance; Multi-compartment vehicle routing problem; Petrol distribution"),
                ("研究背景：", "成品油配送是石油供应链中的关键环节，现有研究较少同时考虑燃油成本对车队组成决策的影响，且缺乏量化不同车型经济行驶距离范围的有效方法。"),
                ("研究方法：", "提出经济行驶距离（Economical Traveling Distance, ETD）方法，通过计算不同车型的燃料消耗率，量化各车型的最佳行驶距离范围。在此基础上，开发了基于ETD的启发式算法。"),
                ("模型介绍：", "研究将问题建模为考虑燃油成本的车队规模与混合车辆路径问题的新变体。模型核心包含三个相互关联的决策层面：①车队组成决策——确定各类型车辆的使用数量和组合；②路径规划决策——确定每辆车访问加油站的顺序和路线；③装载决策——确定各油品如何分配到车辆的多隔舱中。算法模块包括需求分类策略、车型调整策略、装载整合策略和路线整合策略。"),
                ("实验设计：", "以中国大连市某大型成品油配送企业为案例，基于百度地图获取了110个加油站的实际地理数据，结合一个配送中心和多种车型的真实运营参数进行实验验证。"),
            ],

            "conclusion": [
                "研究得出以下主要结论：",
                "①ETD方法能够快速有效地计算出各类型车辆的经济行驶距离范围。大型车辆适合长途运输，小型车辆适合短途配送。",
                "②基于ETD的车队分配算法能够显著降低运输企业的燃油成本。",
                "③该方法不仅适用于当前运营决策，还可为企业的长期车队配置规划提供量化依据。",
                "④ETD方法具有良好的通用性，可推广至其他使用重型混合车队的物流领域。",
            ],

            "harvest": "论文在建模时将车队组成、路径规划、装载决策作为一个联系的整体建模的思路值得学习借鉴。我的毕业论文中也涉及到混装和缺载，在考虑多隔舱时「决策分层+整体优化」的思路可以作为一种参考。另外，论文对于多隔舱的处理方式，允许一个隔舱服务多个加油站，与我的课题中的装载场景有一定的相关性。",
        },
    }

    generate_report(content)
